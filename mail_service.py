import smtplib
import logging
import os
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from datetime import datetime
import io
import re

logger = logging.getLogger(__name__)

class MailService:
    """Service class for handling email operations"""
    
    def __init__(self, settings_provider):
        """
        Initialize with a settings provider
        
        Args:
            settings_provider: An object with a get_value method to retrieve settings
        """
        self.settings = settings_provider
    
    def _get_smtp_settings(self):
        """Get SMTP settings from the settings provider"""
        return {
            'server': self.settings.get_value('SMTP_SERVER'),
            'port': self.settings.get_value('SMTP_PORT', '587'),
            'username': self.settings.get_value('SMTP_USERNAME'),
            'password': self.settings.get_value('SMTP_PASSWORD'),
            'use_tls': self.settings.get_value('SMTP_USE_TLS', 'True').lower() == 'true',
            'from_addr': self.settings.get_value('MAIL_FROM')
        }
    
    def validate_settings(self):
        """
        Validate that required SMTP settings are configured
        
        Returns:
            tuple: (is_valid, error_message)
        """
        smtp_settings = self._get_smtp_settings()
        
        if not smtp_settings['server']:
            return False, "SMTP server is not configured"
            
        if not smtp_settings['from_addr']:
            return False, "From email address is not configured"
            
        return True, None
    
    def _get_replacement_dict(self, contact):
        """Get dictionary of replacements for the template document"""
        return {
            '#first_name': contact.first_name or '',
            '#last_name': contact.last_name or '',
            '#name': f"{contact.first_name or ''} {contact.last_name or ''}".strip(),
            '#NAME': f"{contact.first_name or ''} {contact.last_name or ''}".strip().upper(),
            '#pin': contact.pin or '',
            '#PIN': contact.pin or '',
            '#date': datetime.now().strftime('%d/%m/%Y'),
            '#department': contact.department or '',
            '#dep': contact.department or '',
            '#affil': contact.title or '',
            '#city': 'Ξάνθη'  # Default city, can be adjusted as needed
        }
    
    def _replace_text_in_document(self, doc, replacements):
        """Replace placeholders while preserving formatting in the document"""
        # For each paragraph in the document
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, replacements)

        # For each table in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, replacements)

    def _process_paragraph(self, paragraph, replacements):
        """Process a single paragraph for replacements while preserving formatting"""
        # Skip empty paragraphs
        if not paragraph.text:
            return

        # Check if paragraph contains any replacements before proceeding
        has_replacements = any(key in paragraph.text for key in replacements.keys())
        if not has_replacements:
            return
            
        # Store original runs with their formatting
        runs_with_text = []
        for run in paragraph.runs:
            runs_with_text.append((run, run.text))
            
        # Create a map of where each run's text appears in the paragraph text
        paragraph_text = paragraph.text
        run_positions = []
        position = 0
        
        for run, text in runs_with_text:
            if text:
                run_start = paragraph_text.find(text, position)
                if run_start >= 0:
                    run_positions.append((run, run_start, run_start + len(text)))
                    position = run_start + len(text)
        
        # Apply replacements to the full paragraph text
        new_text = paragraph_text
        for key, value in replacements.items():
            if key in new_text:
                new_text = new_text.replace(key, str(value))
                
        # If text didn't change, no need to continue
        if new_text == paragraph_text:
            return
            
        # Clear all runs first
        for run, _ in runs_with_text:
            run.text = ""
            
        # If we have run position data, try to preserve formatting
        if run_positions:
            # Sort by start position
            run_positions.sort(key=lambda x: x[1])
            
            # Apply text to runs while trying to maintain formatting
            last_pos = 0
            for run, start, end in run_positions:
                # Calculate proportional positions in the new text
                proportion_start = start / len(paragraph_text) if paragraph_text else 0
                proportion_end = end / len(paragraph_text) if paragraph_text else 1
                
                new_start = int(proportion_start * len(new_text))
                new_end = int(proportion_end * len(new_text))
                
                # Adjust for boundaries
                new_start = max(last_pos, min(new_start, len(new_text)))
                new_end = max(new_start, min(new_end, len(new_text)))
                
                # Apply the corresponding slice of new text to this run
                if new_start < new_end:
                    run.text = new_text[new_start:new_end]
                    last_pos = new_end
            
            # If there's remaining text, add it to the last run
            if last_pos < len(new_text) and runs_with_text:
                runs_with_text[-1][0].text += new_text[last_pos:]
        else:
            # Fallback: put all text in the first run
            if runs_with_text:
                runs_with_text[0][0].text = new_text
    
    def _create_pdf_document(self, contact):
        """Create PDF document from DOCX template with contact's PIN information"""
        try:
            # Try to import necessary libraries
            from docx import Document
            from docx2pdf import convert
            import os
            
            # Use the template file from docs folder
            template_path = os.path.join('docs', 'template.docx')
            if os.path.exists(template_path):
                # Use the template file
                doc = Document(template_path)
                
                # Replace placeholders in the document
                replacements = self._get_replacement_dict(contact)
                self._replace_text_in_document(doc, replacements)
                
                # Create temp directory if it doesn't exist
                temp_dir = "temp"
                os.makedirs(temp_dir, exist_ok=True)
                
                # Save the modified DOCX to a temporary file
                contact_id_safe = str(contact.id) if contact.id else "unknown"
                temp_docx = os.path.join(temp_dir, f"temp_{contact_id_safe}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx")
                temp_pdf = os.path.join(temp_dir, f"temp_{contact_id_safe}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf")
                
                doc.save(temp_docx)
                
                try:
                    # Convert DOCX to PDF using docx2pdf
                    convert(temp_docx, temp_pdf)
                    
                    # Read the PDF file
                    with open(temp_pdf, 'rb') as pdf_file:
                        pdf_bytes = pdf_file.read()
                    
                    # Clean up temporary files
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                    if os.path.exists(temp_pdf):
                        os.remove(temp_pdf)
                        
                    return pdf_bytes
                    
                except Exception as conversion_error:
                    logger.error(f"Error converting with docx2pdf: {str(conversion_error)}")
                    # Attempt fallback to PyMuPDF if docx2pdf fails
                    import fitz
                    logger.info("Attempting fallback to PyMuPDF conversion")
                    doc = fitz.open(temp_docx)
                    pdf_bytes = doc.convert_to_pdf()
                    doc.close()
                    
                    # Clean up temp file
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                    
                    return pdf_bytes
            else:
                # Handle case where template is missing
                logger.error("DOCX template not found. Cannot create PDF.")
                raise FileNotFoundError("DOCX template not found at path: " + template_path)
        except ImportError as e:
            logger.error(f"Import error: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error creating PDF from DOCX: {str(e)}")
            raise
    
    def send_pin_email(self, contact):
        """
        Send PIN email to a contact with PDF attachment from template.docx
        
        Args:
            contact: Contact object with email, first_name, last_name, and pin attributes
            
        Returns:
            tuple: (success, error_message)
        """
        if not contact.email or not contact.pin:
            return False, "Contact must have both email and PIN set"
        
        # Get SMTP settings
        smtp_settings = self._get_smtp_settings()
        
        # Validate SMTP settings
        is_valid, error = self.validate_settings()
        if not is_valid:
            return False, error
        
        try:
            # Create the message
            msg = MIMEMultipart()
            msg['From'] = smtp_settings['from_addr']
            msg['To'] = contact.email
            msg['Subject'] = self.settings.get_value('MAIL_SUBJECT', 'Your PIN Document')
            
            # Add simple text body
            body = "Please find your PIN document attached."
            msg.attach(MIMEText(body, 'plain'))
            
            try:
                # Create and attach PDF document
                pdf_bytes = self._create_pdf_document(contact)
                
                # Properly attach the PDF with correct MIME type
                attachment = MIMEApplication(pdf_bytes, _subtype="pdf")
                attachment.add_header('Content-Disposition', 'attachment', 
                                      filename=f"PIN-{contact.last_name}.pdf")
                msg.attach(attachment)
            except Exception as pdf_error:
                logger.error(f"Failed to create PDF attachment: {str(pdf_error)}")
                return False, f"Failed to create PDF attachment: {str(pdf_error)}"
            
            # Send email
            with smtplib.SMTP(smtp_settings['server'], int(smtp_settings['port'])) as server:
                if smtp_settings['use_tls']:
                    server.starttls()
                if smtp_settings['username'] and smtp_settings['password']:
                    server.login(smtp_settings['username'], smtp_settings['password'])
                server.send_message(msg)
                
            return True, None
        except Exception as e:
            logger.error(f"Error sending PIN email: {str(e)}")
            return False, str(e)
